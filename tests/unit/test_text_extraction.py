"""Tests for text extraction service."""

import io

import pytest
from app.core.exceptions import DocumentProcessingError
from app.services.text_extraction import TextExtractor
from docx import Document
from PyPDF2 import PdfWriter


def create_test_pdf(text_content: str) -> bytes:
    """Create a simple test PDF file.

    Args:
        text_content: Text to include in the PDF

    Returns:
        PDF file as bytes
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, text_content)
    c.save()
    buffer.seek(0)
    return buffer.read()


def create_test_docx(paragraphs: list[str]) -> bytes:
    """Create a test DOCX file.

    Args:
        paragraphs: List of paragraph texts

    Returns:
        DOCX file as bytes
    """
    doc = Document()
    for para_text in paragraphs:
        doc.add_paragraph(para_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def test_extract_from_pdf_success():
    """Test successful PDF text extraction."""
    test_text = "This is a test PDF document."
    pdf_content = create_test_pdf(test_text)

    extractor = TextExtractor()
    result = extractor.extract_from_pdf(pdf_content)

    assert test_text in result
    assert len(result) > 0


def test_extract_from_pdf_empty():
    """Test PDF extraction with empty PDF."""
    # Create empty PDF with PdfWriter
    buffer = io.BytesIO()
    writer = PdfWriter()
    writer.write(buffer)
    buffer.seek(0)
    empty_pdf = buffer.read()

    extractor = TextExtractor()

    with pytest.raises(DocumentProcessingError, match="PDF file has no pages"):
        extractor.extract_from_pdf(empty_pdf)


def test_extract_from_pdf_invalid():
    """Test PDF extraction with invalid PDF data."""
    invalid_pdf = b"This is not a PDF file"

    extractor = TextExtractor()

    with pytest.raises(
        DocumentProcessingError, match="Failed to extract text from PDF"
    ):
        extractor.extract_from_pdf(invalid_pdf)


def test_extract_from_docx_success():
    """Test successful DOCX text extraction."""
    test_paragraphs = [
        "This is the first paragraph.",
        "This is the second paragraph.",
        "This is the third paragraph.",
    ]
    docx_content = create_test_docx(test_paragraphs)

    extractor = TextExtractor()
    result = extractor.extract_from_docx(docx_content)

    for para in test_paragraphs:
        assert para in result
    assert len(result) > 0


def test_extract_from_docx_with_tables():
    """Test DOCX extraction with tables."""
    doc = Document()
    doc.add_paragraph("Document with table")

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Cell 1"
    table.rows[1].cells[1].text = "Cell 2"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_content = buffer.read()

    extractor = TextExtractor()
    result = extractor.extract_from_docx(docx_content)

    assert "Document with table" in result
    assert "Header 1" in result
    assert "Header 2" in result
    assert "Cell 1" in result
    assert "Cell 2" in result


def test_extract_from_docx_empty():
    """Test DOCX extraction with empty document."""
    doc = Document()
    # Don't add any content

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    empty_docx = buffer.read()

    extractor = TextExtractor()

    with pytest.raises(DocumentProcessingError, match="DOCX file has no paragraphs"):
        extractor.extract_from_docx(empty_docx)


def test_extract_from_docx_invalid():
    """Test DOCX extraction with invalid DOCX data."""
    invalid_docx = b"This is not a DOCX file"

    extractor = TextExtractor()

    with pytest.raises(
        DocumentProcessingError, match="Failed to extract text from DOCX"
    ):
        extractor.extract_from_docx(invalid_docx)


def test_extract_from_txt_utf8():
    """Test TXT extraction with UTF-8 encoding."""
    test_text = "This is a test text file with UTF-8 encoding. ™ © ®"
    txt_content = test_text.encode("utf-8")

    extractor = TextExtractor()
    result = extractor.extract_from_txt(txt_content)

    assert result == test_text
    assert len(result) > 0


def test_extract_from_txt_latin1():
    """Test TXT extraction with Latin-1 encoding."""
    test_text = "This is a test text file with special chars: café, naïve"
    txt_content = test_text.encode("latin-1")

    extractor = TextExtractor()
    result = extractor.extract_from_txt(txt_content)

    assert "special chars" in result
    assert len(result) > 0


def test_extract_from_txt_utf8_bom():
    """Test TXT extraction with UTF-8 BOM."""
    test_text = "This is a test text file with UTF-8 BOM."
    txt_content = test_text.encode("utf-8-sig")

    extractor = TextExtractor()
    result = extractor.extract_from_txt(txt_content)

    assert test_text in result or result.strip() == test_text.strip()


def test_extract_from_txt_empty():
    """Test TXT extraction with empty file."""
    empty_txt = b""

    extractor = TextExtractor()

    with pytest.raises(DocumentProcessingError, match="TXT file is empty"):
        extractor.extract_from_txt(empty_txt)


def test_extract_from_txt_whitespace_only():
    """Test TXT extraction with whitespace-only file."""
    whitespace_txt = b"   \n\n\t\t   "

    extractor = TextExtractor()

    with pytest.raises(DocumentProcessingError, match="TXT file is empty"):
        extractor.extract_from_txt(whitespace_txt)


def test_extract_text_pdf():
    """Test extract_text method with PDF file type."""
    test_text = "Testing the extract_text method with PDF."
    pdf_content = create_test_pdf(test_text)

    extractor = TextExtractor()
    result = extractor.extract_text(pdf_content, "pdf")

    assert test_text in result


def test_extract_text_docx():
    """Test extract_text method with DOCX file type."""
    test_paragraphs = ["Testing the extract_text method with DOCX."]
    docx_content = create_test_docx(test_paragraphs)

    extractor = TextExtractor()
    result = extractor.extract_text(docx_content, "docx")

    assert test_paragraphs[0] in result


def test_extract_text_txt():
    """Test extract_text method with TXT file type."""
    test_text = "Testing the extract_text method with TXT."
    txt_content = test_text.encode("utf-8")

    extractor = TextExtractor()
    result = extractor.extract_text(txt_content, "txt")

    assert result == test_text


def test_extract_text_case_insensitive():
    """Test extract_text method with uppercase file type."""
    test_text = "Testing case insensitivity."
    txt_content = test_text.encode("utf-8")

    extractor = TextExtractor()
    result = extractor.extract_text(txt_content, "TXT")

    assert result == test_text


def test_extract_text_unsupported_type():
    """Test extract_text method with unsupported file type."""
    content = b"some content"

    extractor = TextExtractor()

    with pytest.raises(DocumentProcessingError, match="Unsupported file type"):
        extractor.extract_text(content, "exe")


def test_extract_text_empty_file_type():
    """Test extract_text method with empty file type."""
    content = b"some content"

    extractor = TextExtractor()

    with pytest.raises(DocumentProcessingError, match="Unsupported file type"):
        extractor.extract_text(content, "")
