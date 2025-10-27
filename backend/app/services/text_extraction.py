"""Text extraction service for PDF, DOCX, and TXT files."""

import io

import docx
from PyPDF2 import PdfReader

from app.core.exceptions import DocumentProcessingError
from app.core.logging import get_logger

logger = get_logger(__name__)


class TextExtractor:
    """Service for extracting text from various document formats."""

    @staticmethod
    def extract_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file.

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text content

        Raises:
            DocumentProcessingError: If PDF extraction fails
        """
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)

            if len(reader.pages) == 0:
                raise DocumentProcessingError("PDF file has no pages")

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue

            if not text_parts:
                raise DocumentProcessingError("No text could be extracted from PDF")

            extracted_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from PDF")
            return extracted_text

        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}", exc_info=True)
            raise DocumentProcessingError(
                f"Failed to extract text from PDF: {e}"
            ) from e

    @staticmethod
    def extract_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file.

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text content

        Raises:
            DocumentProcessingError: If DOCX extraction fails
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)

            if not doc.paragraphs:
                raise DocumentProcessingError("DOCX file has no paragraphs")

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if not text_parts:
                raise DocumentProcessingError("No text could be extracted from DOCX")

            extracted_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text

        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}", exc_info=True)
            raise DocumentProcessingError(
                f"Failed to extract text from DOCX: {e}"
            ) from e

    @staticmethod
    def extract_from_txt(file_content: bytes) -> str:
        """Extract text from TXT file.

        Args:
            file_content: TXT file content as bytes

        Returns:
            Extracted text content

        Raises:
            DocumentProcessingError: If TXT extraction fails
        """
        try:
            # Try UTF-8 first, then fall back to other encodings
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    # Remove null bytes and excessive whitespace
                    text = text.replace("\x00", "").strip()

                    if not text:
                        raise DocumentProcessingError("TXT file is empty")

                    logger.info(
                        f"Extracted {len(text)} characters from TXT using {encoding}"
                    )
                    return text

                except UnicodeDecodeError as e:
                    if encoding == encodings[-1]:
                        # Last encoding failed, raise error
                        raise DocumentProcessingError(
                            "Failed to decode text file with any supported encoding"
                        ) from e
                    continue

            # Should not reach here, but just in case
            raise DocumentProcessingError("Failed to extract text from TXT file")

        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}", exc_info=True)
            raise DocumentProcessingError(
                f"Failed to extract text from TXT: {e}"
            ) from e

    @classmethod
    def extract_text(cls, file_content: bytes, file_type: str) -> str:
        """Extract text from file based on file type.

        Args:
            file_content: File content as bytes
            file_type: File type (pdf, docx, txt)

        Returns:
            Extracted text content

        Raises:
            DocumentProcessingError: If extraction fails or file type is unsupported
        """
        file_type = file_type.lower()

        if file_type == "pdf":
            return cls.extract_from_pdf(file_content)
        elif file_type == "docx":
            return cls.extract_from_docx(file_content)
        elif file_type == "txt":
            return cls.extract_from_txt(file_content)
        else:
            raise DocumentProcessingError(
                f"Unsupported file type: {file_type}. Supported types: pdf, docx, txt"
            )
